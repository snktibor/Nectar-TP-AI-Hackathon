"""Document text extraction for PDF and DOCX files.

Uses pypdf for PDF and python-docx for DOCX. Returns structured text
with page/section boundaries preserved for downstream chunking.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class PageContent:
    """Text content of a single page or section with its ordinal."""

    page_number: int
    text: str


@dataclass
class ParsedDocument:
    """Full extraction result for one uploaded file."""

    filename: str
    pages: list[PageContent] = field(default_factory=list)
    total_characters: int = 0
    error: str | None = None

    @property
    def page_count(self) -> int:
        return len(self.pages)

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)

    def first_n_pages_text(self, n: int = 2) -> str:
        return "\n\n".join(p.text for p in self.pages[:n])


def extract_pdf(filename: str, content: bytes) -> ParsedDocument:
    """Extract text from a PDF file using pypdf."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        pages: list[PageContent] = []
        total_chars = 0

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                pages.append(PageContent(page_number=i + 1, text=text))
                total_chars += len(text)

        return ParsedDocument(
            filename=filename,
            pages=pages,
            total_characters=total_chars,
        )
    except Exception as exc:
        logger.exception("PDF extraction failed for %s", filename)
        return ParsedDocument(filename=filename, error=str(exc))


def extract_docx(filename: str, content: bytes) -> ParsedDocument:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        paragraphs: list[str] = []
        current_section: list[str] = []
        pages: list[PageContent] = []
        section_index = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                if current_section:
                    section_index += 1
                    pages.append(
                        PageContent(
                            page_number=section_index,
                            text="\n".join(current_section),
                        )
                    )
                    current_section = []
                continue

            current_section.append(text)
            paragraphs.append(text)

        if current_section:
            section_index += 1
            pages.append(
                PageContent(
                    page_number=section_index,
                    text="\n".join(current_section),
                )
            )

        if not pages and paragraphs:
            pages.append(PageContent(page_number=1, text="\n".join(paragraphs)))

        total_chars = sum(len(p.text) for p in pages)

        return ParsedDocument(
            filename=filename,
            pages=pages,
            total_characters=total_chars,
        )
    except Exception as exc:
        logger.exception("DOCX extraction failed for %s", filename)
        return ParsedDocument(filename=filename, error=str(exc))


def parse_document(filename: str, content: bytes) -> ParsedDocument:
    """Route to the correct parser based on file extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_pdf(filename, content)
    if lower.endswith(".docx"):
        return extract_docx(filename, content)
    return ParsedDocument(
        filename=filename,
        error=f"Unsupported file type: {filename}. Only PDF and DOCX are accepted.",
    )
